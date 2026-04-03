"""
Document text extraction module.
Supports PDF (pdfplumber), DOCX (python-docx), and Image (CLI-direct OCR).
"""

import base64
import io
import os
import logging
import subprocess
import tempfile
import uuid

from .utils import clean_text

logger = logging.getLogger(__name__)

def extract_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using pdfplumber (Lazy Import)."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        raise ValueError(f"Could not extract text from PDF: {str(e)}")


def extract_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx (Lazy Import)."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        raise ValueError(f"Could not extract text from DOCX: {str(e)}")


def extract_image(file_bytes: bytes) -> str:
    """
    Extract text from an image using Tesseract CLI directly.
    Bypasses python-pytesseract and NumPy DLL issues.
    """
    # 1. Create a temporary directory/file for the image and the output
    tmp_dir = tempfile.gettempdir()
    unique_id = str(uuid.uuid4())
    img_path = os.path.join(tmp_dir, f"{unique_id}.png")
    out_path_base = os.path.join(tmp_dir, f"{unique_id}")
    out_path_txt = f"{out_path_base}.txt"

    try:
        # 2. Save image bytes to a temp file (Pure Python)
        with open(img_path, "wb") as f:
            f.write(file_bytes)

        # 3. Determine Tesseract command
        tesseract_cmd = os.getenv("TESSERACT_CMD", "tesseract")
        
        # 4. Run Tesseract CLI directly via subprocess
        # We use shell=False for security, and pass input/output paths
        logger.info(f"Running Tesseract CLI: {tesseract_cmd}")
        process = subprocess.run(
            [tesseract_cmd, img_path, out_path_base],
            capture_output=True,
            text=True,
            check=False
        )

        if process.returncode != 0:
            err_msg = process.stderr.strip() or "Tesseract execution failed"
            logger.error(f"Tesseract CLI Error: {err_msg}")
            raise ValueError(f"OCR Error: {err_msg}")

        # 5. Read the resulting text file
        if not os.path.exists(out_path_txt):
            return "No text detected in image."

        with open(out_path_txt, "r", encoding="utf-8") as f:
            text = f.read()

        return text

    except Exception as e:
        logger.error(f"Image extraction via CLI failed: {e}")
        # Check for specific DLL block signature
        if "DLL load failed" in str(e) or "_umath_linalg" in str(e):
             raise ValueError("OCR System Blocked: The system's application control policy is blocking low-level AI libraries.")
        raise ValueError(f"Image processing failed: {str(e)}")

    finally:
        # 6. Cleanup temp files
        for p in [img_path, out_path_txt]:
            if os.path.exists(p):
                try: os.remove(p)
                except: pass


def extract_text(file_type: str, file_base64: str) -> str:
    """
    Main dispatcher: decode base64, route to the correct extractor,
    and return cleaned text.
    """
    try:
        if "," in file_base64 and file_base64.startswith("data:"):
            file_base64 = file_base64.split(",", 1)[1]
        file_bytes = base64.b64decode(file_base64)
    except Exception as e:
        raise ValueError(f"Invalid base64 data: {str(e)}")

    file_type = file_type.lower().strip()

    extractors = {
        "pdf": extract_pdf,
        "docx": extract_docx,
        "image": extract_image,
        "png": extract_image,
        "jpg": extract_image,
        "jpeg": extract_image,
        "tiff": extract_image,
        "bmp": extract_image,
        "webp": extract_image,
    }

    extractor_fn = extractors.get(file_type)
    if not extractor_fn:
        raise ValueError(f"Unsupported file type: {file_type}.")

    raw_text = extractor_fn(file_bytes)

    if not raw_text or not raw_text.strip():
        raise ValueError("No text could be extracted from the document.")

    return clean_text(raw_text)
